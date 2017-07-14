from docx import Document
import io
import re
import datetime


def match_price(paragraph, i):
    match = re.match("(\d+)(.*)?(Euro|€)", paragraph)
    if match is not None:
        return {"price": int(match.group(1)), "paragraph": i}


def update_document_contents(document, details):
    f = io.BytesIO(document)
    doc = Document(f)

    # Search for all pricings
    pricings = {}
    for j in range(0, len(doc.paragraphs)):
        match = re.match("(\d+)(.*)?(Euro|€)", doc.paragraphs[j].text)
        if match is not None:
            price = int(match.group(1))
            if price in pricings:
                pricings[price].append(j)
            else:
                pricings[price] = [j]

    # Find the highest price (could be done in the loop above but whatever)
    highest_price = {"price": 0, "paragraph": []}
    for price in pricings:
        if price > highest_price['price']:
            highest_price = {"price": price, "paragraph": pricings[price]}

    # Find the second highest price
    second_highest_price = {"price": 0, "paragraph": []}
    for price in pricings:
        if highest_price['price'] > price > second_highest_price['price']:
            second_highest_price = {"price": price, "paragraph": pricings[price]}

    # Save the lines (with prices) to highlight
    paragraphs_to_highlight = [highest_price['paragraph'][0], second_highest_price['paragraph'][1]]

    # Set up entries to replace
    mappings = {
        "Name": details['name'],
        "Matrikel": details['matrikelnr'],
        "Telefon": details['phone'] + " || " + details['mail'],
        "Anschrift": details['streetAndCity']
    }

    location_and_date = details['location'] + ', ' + datetime.datetime.now().strftime("%d.%m.%Y")
    for i in range(0, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        l = p.runs
        for j in range(len(l)):
            # Replace underscores
            if re.match("(_+_+)", l[j].text) is not None:
                l[j].text = re.sub("(_+_+)", "", l[j].text)

            # Replace the empty checkbox with a checkmark and
            if "" in l[j].text:
                l[j].text = l[j].text.replace("", "")
            # add an arrow at the end of the line
            if re.match(".*(" + details['quarter'] + ")", l[j].text) is not None:
                l[j].text = l[j].text + " <---"

            # Replace the location and date
            if 'Ort, Datum' in l[j].text:
                l[j].text = l[j].text.replace('Ort, Datum', location_and_date)
            for mapping in mappings:
                match = re.match(".*(" + mapping + "([^:]*)?:)", l[j].text)
                if match is not None:
                    l[j].text = l[j].text.replace(match.group(0), match.group(0) + " " + mappings[mapping])

            # Select the prices we want
            if i in paragraphs_to_highlight:
                l[j].text = l[j].text + " <========="

    doc.add_paragraph(details['name'])
    doc.add_paragraph(details['footnote'])

    # Dirty formatting fix (add a newline in front of 'Anschrift:')
    for p in doc.paragraphs:
        l = p.runs
        for j in range(len(l)):
            if "Anschrift:" in l[j].text:
                l[j].text = l[j].text.replace("Anschrift:", "\nAnschrift:")

    doc.save(f)
    res = f.getvalue()
    f.close()
    return res
